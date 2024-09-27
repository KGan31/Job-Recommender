import os
import sys
import json
import hashlib
import pandas as pd
import google.generativeai as genai
from flask import session, make_response
from bson import json_util
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from datetime import datetime
from flask_cors import CORS, cross_origin
from flask import Flask, request, jsonify, send_file, Response
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches
from docx2pdf import convert
import pythoncom

current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

# Importing models
from model.recommend import recommend_courses, suggest_queries

aspiration_model = SentenceTransformer("distilbert-base-nli-stsb-mean-tokens")
df = pd.read_csv("../assets/data/modified_courses.csv")

GEMINI_API_KEY = "AIzaSyD7c4ZO6Y91WpU7VxOrjtejItUTrmYxScM"
# GOOGLE_GEMINI_ENDPOINT = 'https://gemini.googleapis.com/v1beta/'
genai.configure(api_key=GEMINI_API_KEY)
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "application/json",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
    # safety_settings = Adjust safety settings
    # See https://ai.google.dev/gemini-api/docs/safety-settings
)

id = ""

app = Flask(__name__)

app.secret_key = b"0392d5f96b458978e5597c44daaa3817bc0124424e2eab1e9024e182bade4b28"
app.config.update(SESSION_COOKIE_SAMESITE="None", SESSION_COOKIE_SECURE=True)

cors = CORS(app, resources={r"/*": {"origins": "http://localhost/*"}})

uri = "mongodb+srv://vedantnimjed:fRycEHrkwwWDlrQM@cluster0.jmmjy.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"

client = MongoClient(uri, server_api=ServerApi("1"))

# Send a ping to confirm a successful connection
try:
    client.admin.command("ping")
    print("Successfully connected to MongoDB!")
except Exception as e:
    print(e)

db = client.job_recommend


@app.route("/api/register", methods=["POST"])
def register():
    json_string = request.get_data(as_text=True)
    user_details = json.loads(json_string)
    print(user_details["email"])

    if db.Users.find_one({"email": user_details["email"]}):
        res = {"message": "Username already exists"}
        response = app.response_class(
            response=json.dumps(res), status=400, mimetype="application/json"
        )
        return response

    new_user = {
        "email": user_details["email"],
        "name": user_details["name"],
        "password": hashlib.sha256(
            user_details["password"].encode("utf-8")
        ).hexdigest(),
        "profile_score": 0,
        "github": user_details["github"] if "github" in user_details else "",
        "linkedIn": user_details["linkedIn"] if "linkedIn" in user_details else "",
        "education": user_details["education"] if "education" in user_details else [],
        "skills": user_details["skills"] if "skills" in user_details else [],
        "experiences": user_details["work_ex"] if "work_ex" in user_details else [],
        "projects": user_details["projects"] if "projects" in user_details else [],
        "extracurriculars": (
            user_details["extracurriculars"]
            if "extracurriculars" in user_details
            else []
        ),
    }

    _id = db.Users.insert_one(new_user)
    # session['id'] = _id.inserted_id

    response = app.response_class(
        response=json.dumps({"message": "Registered successfully"}),
        status=200,
        mimetype="application/json",
    )
    return response


@app.route("/api/login", methods=["POST"])
def login():
    user_details = json.loads(request.get_data())
    print(user_details)

    if not user_details["email"] or not user_details["password"]:
        return Response(
            json.dumps({"message": "Username and password are required"}), status=400
        )

    user = db.Users.find_one({"email": user_details["email"]})

    if (
        user
        and hashlib.sha256(user_details["password"].encode("utf-8")).hexdigest()
        == user["password"]
    ):
        response = app.response_class(
            response=json.dumps(json.dumps({"message": "Login successful"})),
            status=200,
            mimetype="application/json",
        )
        # session['id'] = str(user["_id"])
        # print(session['id'])
        global id
        id = user["email"]
        print(id)
        return response
    else:
        # return Response(, status=400)
        response = app.response_class(
            response=json.dumps({"message": "Incorrect user name or password"}),
            status=400,
            mimetype="application/json",
        )
        return response


@app.route("/api/jobs", methods=["GET"])
def list_jobs():
    jobs = json.loads(json_util.dumps(db.Jobs.find({})))

    def sorter(job):
        try:
            # pprint(job)
            return len(set(job["skills_req"]).intersection(set(user_skills)))
        except:
            return 0

    user_skills = db.Users.find_one({"email": id})["skills"]

    jobs.sort(key=sorter, reverse=True)

    print(user_skills)
    print([job["skills_req"] for job in jobs])

    return jsonify({"jobs": jobs, "user_skills": user_skills})


@app.route("/api/profile/<string:email>", methods=["GET"])
def get_user_profile(email):
    # Query the database for the user with the specified email
    user = db.Users.find_one({"email": email})  # Exclude sensitive data like password

    if user:
        return json.loads(json_util.dumps(user)), 200
    else:
        return jsonify({"error": "User not found"}), 404
    
def add_user_profile(profile_info):
    project_descriptions = [
        proj.get("description", "") for proj in profile_info["projects"]
    ]
    work_descriptions = [
        work.get("job_description", "") for work in profile_info["experiences"]
    ]
    combined_data = (
        "Projects: "
        + " ".join(project_descriptions)
        + " Work Experience: "
        + " ".join(work_descriptions)
    )
    prompt = (
        combined_data
        + " "
        + "\n"
        + "Give a list of core skills used in these projects. Do not go overboard with it. Output in format like {skills: [python, cpp]}"
    )
    response = model.generate_content(prompt)
    res = json.loads(response.text)
    skills = list(set(profile_info["skills"] + res["skills"]))
    print(skills)

    db.Users.update_one(
        {"email": id},
        {
            "$set": {
                "github": profile_info["github"],
                "linkedIn": profile_info["linkedIn"],
                "education": profile_info["education"],
                "skills": skills,
                "experiences": profile_info["experiences"],
                "projects": profile_info["projects"],
                "extracurriculars": profile_info["extracurriculars"],
            }
        },
    )
    
@app.route("/api/get-profile-questions", methods=["GET"])
def get_profile_questions():
    data_param = request.args.get('data') # {"skills": [skill1, skill2, ...]}
    if data_param:
        try:
            # Decode and parse the JSON string back into a Python object
            data = json.loads(data_param)
            skills = data.get('skills', [])
            # print(skills)
            skills = str(skills)
            prompt = skills + """\n
Based on the given list of skills, give 10 conceptual, medium-level and domain specific questions in single-correct MCQ format for judging how well a person knows each these skills fundamentally, along with their correct answers. Do not give any explanations, just a list of questions in format.
[
{question: "question", options: ["A":"option1", "B":"option2"...], correct: ["A"]}
]
Make the order of the options less predictable"""
            
            response = model.generate_content(prompt)
            res = json.loads(response.text)

            return res
        except (json.JSONDecodeError, TypeError) as e:
            return jsonify({'error': 'Invalid JSON format'}), 400
    else:
        return jsonify({'error': 'Missing data'}), 400


@app.route("/api/save-profile", methods=["POST"])
def add_user_skills():
    profile_info = json.loads(request.get_data())
    print(profile_info)
    global id
    print(id)

    add_user_profile(profile_info)

    response = app.response_class(
        response=json.dumps({"message": "Skills added"}),
        status=200,
        mimetype="application/json",
    )
    return response

@app.route("/api/update-profile-score", methods=["POST"])
def update_profile_score():
    profile_info = json.loads(request.get_data())
    db.Users.update_one({"email": id}, {"$set": {"profile_score": profile_info['profile_score']}})

    response = app.response_class(
        response=json.dumps({"message": "Updated user profile"}),
        status=200,
        mimetype="application/json",
    )
    return response

@app.route("/api/recommend-course", methods=["GET"])
def getCourses():
    skills = request.args.getlist("skills")
    res = []

    for s in skills:
        c = db.Courses.find({"skill": s.lower()})
        res.append(json.loads(json_util.dumps(c)))

    print(res)

    response = app.response_class(
        response=json.dumps({"recommendation": res}),
        status=200,
        mimetype="application/json",
    )
    return response


@app.route("/api/aspirations", methods=["POST"])
def get_courses_from_aspirations():
    user_aspiration = json.loads(request.get_data())["aspiration"]

    # Get recommendations
    recommended_courses = recommend_courses(user_aspiration, aspiration_model, df)

    # Print recommendations
    return jsonify({"courses": recommended_courses})


@app.route("/api/similar-questions", methods=["POST"])
def get_similar_questions():
    pass
    # TO DO


@app.route("/api/resume", methods=["POST"])
def get_resume():
    data = json.loads(request.get_data())
    resume_details = data['log']
    flag = data['flag'] # flag to check if there were some changes made by the user. If yes then flag will be False
    print(resume_details)
    print(flag)

    # if changes made by the user update the user profile
    if flag == False:
        add_user_profile(resume_details)

    name = resume_details["name"]
    email = resume_details["email"]
    github = resume_details["github"]
    linkedin = resume_details["linkedIn"]
    skills = resume_details["skills"]
    education = resume_details["education"]
    experience = resume_details["experiences"]
    projects = resume_details["projects"]
    extra = resume_details["extracurriculars"]

    doc = Document("resume-template.docx")
    for paragraph in doc.paragraphs:
        # Replace name
        if "{{name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{name}}", name)

        if "{{email}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{email}}", email)

        if "{{github}}" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "{{github}} ", f"{github} " if github else ""
            )

        if "{{linkedin}}" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "{{linkedin}} ", f"{linkedin} " if linkedin else ""
            )

        if "{{skills}}" in paragraph.text:
            skills_text = ", ".join(skills)
            paragraph.clear()
            text = paragraph.add_run(skills_text)
            text.bold = False

        if "{{education}}" in paragraph.text:
            paragraph.clear()

            # Iterate through each experience and add it in the required format
            for p in education:
                # Add project title and make it bold
                run_title = paragraph.add_run(
                    f'{p["degree"]} at {p["university"]}, {p["location"]}'
                )
                run_title.bold = True

                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )

                tab = paragraph.add_run("\t")

                # Add spaces and the date range, right-aligned
                run_date = paragraph.add_run(f"{p['from']} - {p['to']}")
                run_date.bold = False

                # Add a line break for the description
                paragraph.add_run("\n")

                # Add the description
                run_desc = paragraph.add_run(f'CGPA: {p["cgpa"]}')
                run_desc.bold = False
                endline = paragraph.add_run("\n")

        # Replace experience
        if "{{projects}}" in paragraph.text:
            # Clear the paragraph text
            paragraph.clear()

            # Iterate through each experience and add it in the required format
            for p in projects:
                # Add project title and make it bold
                run_title = paragraph.add_run(p["title"])
                run_title.bold = True

                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )

                tab = paragraph.add_run("\t")

                # Add spaces and the date range, right-aligned
                run_date = paragraph.add_run(f"{p['from']} - {p['to']}")
                run_date.bold = False

                # Add a line break for the description
                paragraph.add_run("\n")

                # Add the description
                run_desc = paragraph.add_run(p["description"])
                run_desc.bold = False
                endline = paragraph.add_run("\n")

        if "{{experience}}" in paragraph.text:
            # Clear the paragraph text
            paragraph.clear()

            # Iterate through each experience and add it in the required format
            for p in experience:
                # Add project title and make it bold
                run_title = paragraph.add_run(p["position"])
                run_title.bold = True

                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(7.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )

                tab = paragraph.add_run("\t")

                # Add spaces and the date range, right-aligned
                run_date = paragraph.add_run(f"{p['from']} - {p['to']}")
                run_date.bold = False

                # Add a line break for the description
                paragraph.add_run("\n")

                # Add the description
                run_desc = paragraph.add_run(p["job_description"])
                run_desc.bold = False

                endline = paragraph.add_run("\n")

        if "{{extra}}" in paragraph.text:
            # Clear the paragraph text
            paragraph.clear()

            # Iterate through each experience and add it in the required format
            for p in extra:
                # Add project title and make it bold
                run_title = paragraph.add_run(p)
                run_title.bold = False
                endline = paragraph.add_run("\n")

    docx_filename = "RESUME.docx"
    pdf_filename = "RESUME.pdf"
    doc.save(docx_filename)

    # Ensure COM initialization before converting DOCX to PDF
    pythoncom.CoInitialize()  # Initialize COM library for the current thread
    try:
        convert(docx_filename, pdf_filename)  # Convert DOCX to PDF
    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM library to clean up resources

    # Remove the DOCX file if you don't need it anymore
    os.remove(docx_filename)

    response = make_response(send_file(pdf_filename, as_attachment=True))
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = 'inline; filename="RESUME.pdf"'
    return response


if __name__ == "__main__":
    app.run(debug=True)
